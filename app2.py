import os
import streamlit as st
from pdf2image import convert_from_path
import easyocr
from google.cloud import translate_v2 as translate
from google.oauth2 import service_account
import numpy as np
import torch
import tempfile
import gc
from langdetect import detect, DetectorFactory
import logging
from PIL import Image, ImageEnhance
import io
import json
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Set a seed for reproducibility in language detection
DetectorFactory.seed = 0

# Set page configuration
st.set_page_config(page_title="Bulk OCR & Translation App", layout="wide")

# Custom CSS for a softer UI with a bold color scheme


def apply_custom_css():
    st.markdown("""
    <style>
    .stApp {
        max-width: 1200px;
        margin: 0 auto;
        background-color: #121212;
        color: #ffffff;
    }
    .stButton > button {
        background-color: #333333;
        color: white;
        border-radius: 20px;
        padding: 10px 20px;
        font-weight: bold;
    }
    .stSelectbox, .stTextArea {
        border-radius: 10px;
        border: 2px solid #333333;
    }
    .output-box {
        border: 2px solid #333333;
        border-radius: 10px;
        padding: 20px;
        margin-bottom: 20px;
        background-color: #202020;
        color: #ffffff;
    }
    h1, h2, h3 {
        color: #ffffff;
    }
    </style>
    """, unsafe_allow_html=True)


# Function to set GPU memory strategy


def set_gpu_memory_strategy():
    os.environ['PYTORCH_CUDA_ALLOC_CONF'] = 'expandable_segments:True'
    if torch.cuda.is_available():
        torch.cuda.set_per_process_memory_fraction(0.7)

# Function to enhance image for better OCR


def enhance_image(image):
    enhancer = ImageEnhance.Contrast(image)
    enhanced_image = enhancer.enhance(1.5)  # Increase contrast
    return enhanced_image

# Function to process each page


def process_page(page, reader):
    try:
        enhanced_page = enhance_image(page)
        img_array = np.array(enhanced_page.convert('RGB'))
        result = reader.readtext(img_array, detail=0, paragraph=True)
        # Join paragraphs with double newlines
        extracted_text = "\n\n".join(result)
        return extracted_text
    except Exception as e:
        logger.error(f"Error processing page: {str(e)}")
        return None

# Updated translate function using Google Cloud API


def translate_text(text, target_language, credentials_json):
    try:
        credentials_dict = json.loads(credentials_json)
        credentials = service_account.Credentials.from_service_account_info(
            credentials_dict)
        client = translate.Client(credentials=credentials)
        translated_text = client.translate(
            text, target_language=target_language)
        return translated_text['translatedText']
    except Exception as e:
        logger.error(f"Translation error: {str(e)}")
        return f"Translation error: {str(e)}"

# Function to format text with basic structure preservation


def format_text(text):
    lines = text.split('\n')
    formatted_text = []

    for line in lines:
        stripped_line = line.strip()
        if stripped_line:
            if stripped_line.isupper() and len(stripped_line.split()) <= 5:
                formatted_text.append(f"\n# {stripped_line}\n")
            elif stripped_line[0].isupper() and len(stripped_line.split()) <= 10:
                formatted_text.append(f"\n## {stripped_line}\n")
            else:
                formatted_text.append(line)

    return "\n".join(formatted_text)

# Function to save as DOCX


def save_as_docx(pages_text, filename):
    doc = Document()
    for i, page_text in enumerate(pages_text, 1):
        doc.add_heading(f'Page {i}', level=1)
        for line in page_text.split('\n'):
            if line.startswith('# '):
                doc.add_heading(line[2:], level=2)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=3)
            else:
                doc.add_paragraph(line)
        if i < len(pages_text):
            doc.add_page_break()
    doc.save(filename)

# Function to save as PDF


def save_as_pdf(pages_text, filename):
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    styles.add(ParagraphStyle(name='Heading1',
               fontSize=18, textColor=colors.darkblue))
    styles.add(ParagraphStyle(name='Heading2',
               fontSize=16, textColor=colors.darkblue))
    styles.add(ParagraphStyle(name='Heading3',
               fontSize=14, textColor=colors.darkblue))
    styles.add(ParagraphStyle(name='BodyText', fontSize=12,
               leading=14, textColor=colors.black))

    for i, page_text in enumerate(pages_text, 1):
        story.append(Paragraph(f'Page {i}', styles['Heading1']))
        story.append(Spacer(1, 12))
        for line in page_text.split('\n'):
            if line.startswith('# '):
                story.append(Paragraph(line[2:], styles['Heading2']))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], styles['Heading3']))
            else:
                story.append(Paragraph(line, styles['BodyText']))
            story.append(Spacer(1, 6))
        if i < len(pages_text):
            story.append(PageBreak())

    doc.build(story)


def main():
    apply_custom_css()

    st.title("Enhanced Bulk OCR & Translation App")

    # Sidebar
    st.sidebar.title("Settings")

    # Service Account JSON input
    credentials_json = st.sidebar.text_area(
        "Enter your Google Cloud Service Account JSON",
        height=150
    )

    # Instructions for creating the Service Account JSON
    with st.sidebar.expander("Instructions for Service Account JSON"):
        st.markdown("""
        1. Visit [Google Cloud Console](https://console.cloud.google.com/).
        2. Navigate to **IAM & Admin** > **Service Accounts**.
        3. Create a new service account or select an existing one.
        4. Generate a new JSON key for the service account.
        5. Copy the entire contents of the JSON file and paste it above.
        6. Ensure the service account has the "Cloud Translation API User" role.
        """)

    if not credentials_json:
        st.warning(
            "Please enter your Google Cloud Service Account JSON to enable translation.")
        return

    # File uploader
    uploaded_file = st.file_uploader("Choose a PDF file", type=["pdf"])

    # Destination language options

    translation_options = {
        "English": "en",
        "Spanish": "es",
        "French": "fr",
        "German": "de",
        "Italian": "it",
        "Japanese": "ja",
        "Korean": "ko",
        "Chinese (Simplified)": "zh-CN",
        "Hindi": "hi",
        "Bengali": "bn",
        "Marathi": "mr",
        "Tamil": "ta",
        "Telugu": "te",
        "Gujarati": "gu",
        "Malayalam": "ml",
        "Punjabi": "pa",
        "Urdu": "ur",
        "Kazakh": "kk",
        "Kyrgyz": "ky",
        "Tajik": "tg",
        "Uzbek": "uz",
        "Uyghur": "ug",
        "Azerbaijani": "az",
        "Turkish": "tr"
    }

    target_language = st.selectbox(
        "Select Target Language", list(translation_options.keys()))

    # OCR language selection
    ocr_languages = ['en', 'es', 'fr', 'de', 'it', 'ja', 'ko',
                     'zh_sim', 'hi', 'bn', 'mr', 'ta', 'te', 'gu', 'ml', 'pa', 'ur',
                     'kk', 'ky', 'tg', 'uz', 'ug', 'az', 'tr']
    selected_ocr_languages = st.multiselect(
        "Select OCR Languages", ocr_languages, default=['en'])

    # Output format selection
    output_format = st.selectbox(
        "Select Output Format", ["TXT", "DOCX", "PDF"])

    if uploaded_file is not None:
        try:
            # Save the uploaded file to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                temp_pdf.write(uploaded_file.read())
                temp_pdf_path = temp_pdf.name

            # Load pages from the temporary PDF file
            pages = convert_from_path(temp_pdf_path, 100)

            # Initialize EasyOCR reader
            use_gpu = torch.cuda.is_available()
            reader = easyocr.Reader(
                selected_ocr_languages, gpu=use_gpu, quantize=True)

            all_extracted_text = []
            all_translated_text = []

            # Progress bar
            progress_bar = st.progress(0)

            # Process pages
            for i, page in enumerate(pages):
                extracted_text = process_page(page, reader)
                if extracted_text:
                    try:
                        # Translate text using Google Cloud API
                        translated_text = translate_text(
                            extracted_text, translation_options[target_language], credentials_json)

                        formatted_text = format_text(translated_text)
                        all_translated_text.append(formatted_text)
                    except Exception as e:
                        st.error(f"Error processing page {i+1}: {str(e)}")
                        logger.error(f"Error processing page {i+1}: {str(e)}")

                # Update progress bar
                progress_bar.progress((i + 1) / len(pages))

            # Combine all translated text into a single output
            combined_output = "\n\n".join(all_translated_text)

            # Format the text
            formatted_output = format_text(combined_output)

            # Display the output text
            st.markdown("### Translated Output")
            for i, page_text in enumerate(all_translated_text, 1):
                with st.expander(f"Page {i}"):
                    st.markdown(page_text)

            # Download button for the output file
            if all_translated_text:
                st.markdown("---")
                st.subheader("Download Translated Document")

                if output_format == "TXT":
                    tmp_file = tempfile.NamedTemporaryFile(
                        delete=False, suffix='.txt')
                    with open(tmp_file.name, 'w', encoding='utf-8') as f:
                        for i, page_text in enumerate(all_translated_text, 1):
                            f.write(f"Page {i}\n\n")
                            f.write(page_text)
                            f.write("\n\n" + "="*50 + "\n\n")
                    with open(tmp_file.name, "rb") as file:
                        st.download_button(
                            label="Download Translated Document (TXT)",
                            data=file,
                            file_name="translated_document.txt",
                            mime="text/plain"
                        )
                elif output_format == "DOCX":
                    tmp_file = tempfile.NamedTemporaryFile(
                        delete=False, suffix='.docx')
                    save_as_docx(all_translated_text, tmp_file.name)
                    with open(tmp_file.name, "rb") as file:
                        st.download_button(
                            label="Download Translated Document (DOCX)",
                            data=file,
                            file_name="translated_document.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                elif output_format == "PDF":
                    tmp_file = tempfile.NamedTemporaryFile(
                        delete=False, suffix='.pdf')
                    save_as_pdf(all_translated_text, tmp_file.name)
                    with open(tmp_file.name, "rb") as file:
                        st.download_button(
                            label="Download Translated Document (PDF)",
                            data=file,
                            file_name="translated_document.pdf",
                            mime="application/pdf"
                        )

                os.unlink(tmp_file.name)  # Delete the temporary file

        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
            logger.error(f"An error occurred: {str(e)}")

        finally:
            # Clean up
            gc.collect()
            if use_gpu:
                torch.cuda.empty_cache()


if __name__ == "__main__":
    main()
